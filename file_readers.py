"""
File reading utilities for PDF, DOCX, XLSX and TXT documents.

Every public function takes a filesystem path and returns plain text (or
structured metadata), so the rest of the application never needs to know
which library was used to parse a given file type.

This module is also responsible for path safety: every incoming path is
resolved and validated against config.BASE_DIR before any I/O happens,
which is what keeps the assistant from reading files outside the folder
it was scoped to (e.g. via "../../etc/passwd").
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List

import fitz  # PyMuPDF
import pandas as pd
from docx import Document

from config import config

logger = logging.getLogger(__name__)


class UnsupportedFileTypeError(Exception):
    """Raised when a file extension has no registered reader."""


class PathAccessError(Exception):
    """Raised when a requested path resolves outside the allowed base directory."""


@dataclass
class FileInfo:
    name: str
    path: str
    extension: str
    size_bytes: int
    modified: str
    is_supported: bool


def resolve_safe_path(user_path: str) -> Path:
    """
    Resolve a user-supplied path against config.BASE_DIR and guarantee it
    doesn't escape that directory. Blocks path traversal via '..' segments,
    absolute paths outside the base dir, and symlink escapes.
    """
    raw = Path(user_path)
    candidate = raw.resolve() if raw.is_absolute() else (config.BASE_DIR / raw).resolve()

    try:
        candidate.relative_to(config.BASE_DIR)
    except ValueError:
        raise PathAccessError(
            f"Access denied: '{user_path}' resolves outside the allowed "
            f"base directory '{config.BASE_DIR}'."
        ) from None
    return candidate


def list_files_in_directory(folder_path: str = ".", recursive: bool = True) -> List[FileInfo]:
    """List files under folder_path (relative to BASE_DIR), with metadata."""
    base = resolve_safe_path(folder_path)
    if not base.exists():
        raise FileNotFoundError(f"Folder not found: {folder_path}")
    if not base.is_dir():
        raise NotADirectoryError(f"Not a directory: {folder_path}")

    pattern = "**/*" if recursive else "*"
    results: List[FileInfo] = []
    for p in sorted(base.glob(pattern)):
        if p.is_file():
            stat = p.stat()
            results.append(
                FileInfo(
                    name=p.name,
                    path=str(p.relative_to(config.BASE_DIR)),
                    extension=p.suffix.lower(),
                    size_bytes=stat.st_size,
                    modified=datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
                    is_supported=p.suffix.lower() in config.SUPPORTED_EXTENSIONS,
                )
            )
    return results


def read_pdf(path: Path) -> str:
    text_parts = []
    with fitz.open(path) as doc:
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text().strip()
            if text:
                text_parts.append(f"[Page {page_num}]\n{text}")
    return "\n\n".join(text_parts)


def read_docx(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_xlsx(path: Path) -> str:
    parts = []
    sheets = pd.read_excel(path, sheet_name=None, engine="openpyxl")
    for sheet_name, df in sheets.items():
        parts.append(f"[Sheet: {sheet_name}]")
        parts.append(df.to_string(index=False, na_rep=""))
    return "\n\n".join(parts)


def read_txt(path: Path) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError("utf-8/latin-1", b"", 0, 1, "Could not decode text file")


_READERS = {
    ".pdf": read_pdf,
    ".docx": read_docx,
    ".xlsx": read_xlsx,
    ".xls": read_xlsx,
    ".txt": read_txt,
}


def read_file(file_path: str) -> str:
    """Read any supported document and return its extracted plain text."""
    path = resolve_safe_path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    if not path.is_file():
        raise IsADirectoryError(f"Not a file: {file_path}")

    ext = path.suffix.lower()
    reader = _READERS.get(ext)
    if reader is None:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'. Supported: {sorted(config.SUPPORTED_EXTENSIONS)}"
        )

    logger.info("Reading %s file: %s", ext, path)
    text = reader(path)
    if not text.strip():
        logger.warning("No extractable text found in %s", path)
    return text


def delete_file(file_path: str) -> None:
    """
    Permanently delete a single file within the allowed base directory.
    Raises the same errors as read_file for missing/invalid paths, so
    callers can handle both with the same except clauses.
    """
    path = resolve_safe_path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    if not path.is_file():
        raise IsADirectoryError(f"Not a file: {file_path}")
    path.unlink()
    logger.info("Deleted file: %s", path)


def delete_all_files(folder_path: str = ".", recursive: bool = True) -> List[str]:
    """
    Delete every supported document under folder_path. Returns the list of
    relative paths that were deleted. Unsupported files (source code,
    config files, anything that isn't PDF/DOCX/XLSX/TXT) are left alone
    even if they happen to live in the same folder — this only ever
    touches documents, never the project itself.
    """
    files = list_files_in_directory(folder_path, recursive)
    deleted: List[str] = []
    for f in files:
        if f.is_supported:
            path = resolve_safe_path(f.path)
            try:
                path.unlink()
                deleted.append(f.path)
                logger.info("Deleted file: %s", path)
            except OSError as exc:
                logger.warning("Could not delete %s: %s", path, exc)
    return deleted


def unique_destination_path(folder_path: str, filename: str) -> Path:
    """
    Given a desired filename, return a path inside folder_path that doesn't
    collide with an existing file — appending ' (1)', ' (2)', etc. as
    needed. Used when a caller uploads a file and chooses "keep both"
    instead of overwriting an existing file with the same name.
    """
    base = resolve_safe_path(folder_path)
    stem = Path(filename).stem
    suffix = Path(filename).suffix
    candidate = base / filename
    counter = 1
    while candidate.exists():
        candidate = base / f"{stem} ({counter}){suffix}"
        counter += 1
    return candidate
