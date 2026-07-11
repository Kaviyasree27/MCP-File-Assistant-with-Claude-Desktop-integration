"""
Shared pytest fixtures.

Each test gets an isolated temporary directory that is patched in as
config.BASE_DIR, so tests never touch the real filesystem outside the
sandbox and can't interfere with each other.
"""

from __future__ import annotations

import sys
from pathlib import Path

import openpyxl
import pytest
from docx import Document

# Make the project root importable when running `pytest` from the repo root.
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from config import config  # noqa: E402


@pytest.fixture()
def sandbox(tmp_path, monkeypatch):
    """Point config.BASE_DIR at a fresh temp directory and return it."""
    monkeypatch.setattr(config, "BASE_DIR", tmp_path.resolve())
    return tmp_path


@pytest.fixture()
def sample_txt(sandbox):
    path = sandbox / "notes.txt"
    path.write_text(
        "Project Phoenix kickoff notes.\n"
        "The budget for Project Phoenix is 50000 dollars.\n"
        "Deadline is set for December 2026.\n",
        encoding="utf-8",
    )
    return path


@pytest.fixture()
def sample_docx(sandbox):
    path = sandbox / "report.docx"
    doc = Document()
    doc.add_paragraph("Quarterly Report")
    doc.add_paragraph("Revenue grew by 20 percent this quarter.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Region"
    table.rows[0].cells[1].text = "Sales"
    doc.save(path)
    return path


@pytest.fixture()
def sample_xlsx(sandbox):
    path = sandbox / "data.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sales"
    ws.append(["Region", "Revenue"])
    ws.append(["North", 1000])
    ws.append(["South", 2000])
    wb.save(path)
    return path


@pytest.fixture()
def nested_folder(sandbox):
    (sandbox / "sub").mkdir()
    (sandbox / "sub" / "inner.txt").write_text("Nested file about Project Phoenix budget.", encoding="utf-8")
    (sandbox / "top.txt").write_text("Top level file about Project Phoenix timeline.", encoding="utf-8")
    return sandbox
